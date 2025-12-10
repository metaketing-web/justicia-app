#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pour générer un document Word vierge avec en-tête Porteo et contenu markdown
"""
import sys
import json
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Couleurs Porteo
COLOR_TITLE = RGBColor(0xBA, 0x8A, 0x52)  # #BA8A52 - Or/Bronze pour les titres
COLOR_TEXT = RGBColor(0x17, 0x23, 0x2E)   # #17232E - Bleu foncé pour le texte

def parse_markdown_line(doc, line):
    """Parse une ligne markdown et l'ajoute au document"""
    line = line.strip()
    
    if not line:
        return False
    
    # Titre niveau 1 (###)
    if line.startswith('### '):
        title = line.replace('### ', '').strip()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_TITLE
        p.space_after = Pt(12)
        p.space_before = Pt(18)
        return True
    
    # Titre niveau 2 (####)
    elif line.startswith('#### '):
        title = line.replace('#### ', '').strip()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_TITLE
        p.space_after = Pt(10)
        p.space_before = Pt(14)
        return True
    
    # Liste à puces (-)
    elif line.startswith('- '):
        text = line.replace('- ', '').strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run('• ')
        run.font.color.rgb = COLOR_TEXT
        add_formatted_text(p, text)
        p.space_after = Pt(6)
        return True
    
    # Liste numérotée
    elif re.match(r'^\d+\.\s', line):
        match = re.match(r'^(\d+)\.\s', line)
        num = match.group(1)
        text = re.sub(r'^\d+\.\s', '', line).strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f'{num}. ')
        run.font.color.rgb = COLOR_TEXT
        add_formatted_text(p, text)
        p.space_after = Pt(6)
        return True
    
    return False

def add_formatted_text(paragraph, text):
    """Ajoute du texte avec formatage markdown (gras, italique)"""
    # Remplacer **texte** par gras et *texte* par italique
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.color.rgb = COLOR_TEXT
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.color.rgb = COLOR_TEXT
        else:
            run = paragraph.add_run(part)
            run.font.color.rgb = COLOR_TEXT

def generate_blank_document(title, content, output_path, header_template_path):
    """Génère un document Word vierge avec en-tête Porteo"""
    
    # Charger le template d'en-tête Porteo
    if os.path.exists(header_template_path):
        doc = Document(header_template_path)
    else:
        doc = Document()
    
    # Ajouter un espace après l'en-tête
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Titre principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_title = p_title.add_run(title)
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_TITLE
    p_title.space_after = Pt(24)
    
    # Parser le contenu markdown
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Essayer de parser comme titre ou liste
        if parse_markdown_line(doc, line):
            i += 1
            continue
        
        # Sinon, accumuler les lignes du paragraphe
        if line:
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line or next_line.startswith(('#', '-')) or re.match(r'^\d+\.\s', next_line):
                    break
                para_lines.append(next_line)
                i += 1
            
            # Créer le paragraphe
            text = ' '.join(para_lines)
            p = doc.add_paragraph()
            add_formatted_text(p, text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p.space_after = Pt(12)
        else:
            i += 1
    
    # Sauvegarder
    doc.save(output_path)
    print(f"[BLANK] Document vierge généré: {output_path}")

if __name__ == '__main__':
    if len(sys.argv) < 5:
        print("Usage: python3 generate_blank_document.py <title> <content> <output_path> <header_template_path>")
        sys.exit(1)
    
    title = sys.argv[1]
    content = sys.argv[2]
    output_path = sys.argv[3]
    header_template_path = sys.argv[4]
    
    generate_blank_document(title, content, output_path, header_template_path)

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pour remplir un template Word avec en-tête Porteo et couleurs personnalisées
"""
import sys
import json
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Couleurs Porteo
COLOR_TITLE = RGBColor(0xBA, 0x8A, 0x52)  # #BA8A52 - Or/Bronze pour les titres
COLOR_TEXT = RGBColor(0x17, 0x23, 0x2E)   # #17232E - Bleu foncé pour le texte

def copy_header_footer(source_doc, target_doc):
    """Copie l'en-tête et le pied de page du document source vers le document cible"""
    # Copier les sections
    for src_section, tgt_section in zip(source_doc.sections, target_doc.sections):
        # Copier l'en-tête
        for src_para in src_section.header.paragraphs:
            tgt_para = tgt_section.header.add_paragraph()
            tgt_para.text = src_para.text
            tgt_para.style = src_para.style
            tgt_para.alignment = src_para.alignment
            
            # Copier le formatage des runs
            for src_run, tgt_run in zip(src_para.runs, tgt_para.runs):
                tgt_run.bold = src_run.bold
                tgt_run.italic = src_run.italic
                tgt_run.underline = src_run.underline
                if src_run.font.size:
                    tgt_run.font.size = src_run.font.size
                if src_run.font.color.rgb:
                    tgt_run.font.color.rgb = src_run.font.color.rgb
        
        # Copier le pied de page
        for src_para in src_section.footer.paragraphs:
            tgt_para = tgt_section.footer.add_paragraph()
            tgt_para.text = src_para.text
            tgt_para.style = src_para.style
            tgt_para.alignment = src_para.alignment
            
            for src_run, tgt_run in zip(src_para.runs, tgt_para.runs):
                tgt_run.bold = src_run.bold
                tgt_run.italic = src_run.italic
                tgt_run.underline = src_run.underline
                if src_run.font.size:
                    tgt_run.font.size = src_run.font.size
                if src_run.font.color.rgb:
                    tgt_run.font.color.rgb = src_run.font.color.rgb

def apply_porteo_colors(paragraph, is_title=False):
    """Applique les couleurs Porteo à un paragraphe"""
    for run in paragraph.runs:
        if is_title:
            run.font.color.rgb = COLOR_TITLE
            run.bold = True
            run.font.size = Pt(14)
        else:
            run.font.color.rgb = COLOR_TEXT
            run.font.size = Pt(11)

def replace_placeholder(paragraph, placeholder, value, is_title=False):
    """Remplace un placeholder dans un paragraphe et applique les couleurs"""
    if placeholder in paragraph.text:
        # Remplacer dans les runs
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(value))
                # Appliquer les couleurs Porteo
                if is_title:
                    run.font.color.rgb = COLOR_TITLE
                    run.bold = True
                    run.font.size = Pt(14)
                else:
                    run.font.color.rgb = COLOR_TEXT
                    run.font.size = Pt(11)

def fill_template_with_porteo(template_path, output_path, data, header_template_path):
    """Remplit un template Word avec les données et ajoute l'en-tête Porteo"""
    
    # Charger le template de base
    doc = Document(template_path)
    
    # Charger le template d'en-tête Porteo
    if os.path.exists(header_template_path):
        header_doc = Document(header_template_path)
        
        # S'assurer qu'il y a au moins une section
        if len(doc.sections) == 0:
            doc.add_section()
        
        # Copier l'en-tête et le pied de page du template Porteo
        try:
            # Vider l'en-tête existant
            for section in doc.sections:
                section.header.paragraphs[0].clear()
                
                # Copier l'en-tête Porteo
                for src_para in header_doc.sections[0].header.paragraphs:
                    if src_para.text.strip():  # Ignorer les paragraphes vides
                        tgt_para = section.header.add_paragraph()
                        tgt_para.text = src_para.text
                        tgt_para.alignment = src_para.alignment
                        
                        # Copier le formatage
                        for i, src_run in enumerate(src_para.runs):
                            if i < len(tgt_para.runs):
                                tgt_run = tgt_para.runs[i]
                            else:
                                tgt_run = tgt_para.add_run()
                            
                            tgt_run.text = src_run.text
                            tgt_run.bold = src_run.bold
                            tgt_run.italic = src_run.italic
                            if src_run.font.size:
                                tgt_run.font.size = src_run.font.size
                            if src_run.font.color.rgb:
                                tgt_run.font.color.rgb = src_run.font.color.rgb
                
                # Copier le pied de page Porteo
                if len(header_doc.sections[0].footer.paragraphs) > 0:
                    section.footer.paragraphs[0].clear()
                    for src_para in header_doc.sections[0].footer.paragraphs:
                        if src_para.text.strip():
                            tgt_para = section.footer.add_paragraph()
                            tgt_para.text = src_para.text
                            tgt_para.alignment = src_para.alignment
        except Exception as e:
            print(f"[WARN] Erreur lors de la copie de l'en-tête: {e}")
    
    # Créer les placeholders à partir des données
    placeholders = {}
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"
        placeholders[placeholder] = value
    
    # Remplacer dans tous les paragraphes avec couleurs Porteo
    for paragraph in doc.paragraphs:
        # Détecter si c'est un titre (style Heading ou texte en gras)
        is_title = (paragraph.style.name.startswith('Heading') or 
                   (len(paragraph.runs) > 0 and paragraph.runs[0].bold))
        
        for placeholder, value in placeholders.items():
            replace_placeholder(paragraph, placeholder, value, is_title)
        
        # Appliquer les couleurs même si pas de placeholder
        if paragraph.text.strip():
            apply_porteo_colors(paragraph, is_title)
    
    # Remplacer dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    is_title = (paragraph.style.name.startswith('Heading') or 
                               (len(paragraph.runs) > 0 and paragraph.runs[0].bold))
                    
                    for placeholder, value in placeholders.items():
                        replace_placeholder(paragraph, placeholder, value, is_title)
                    
                    if paragraph.text.strip():
                        apply_porteo_colors(paragraph, is_title)
    
    # Sauvegarder
    doc.save(output_path)
    print(f"[PORTEO] Document généré avec en-tête Porteo: {output_path}")

if __name__ == '__main__':
    if len(sys.argv) < 5:
        print("Usage: python3 fill_template_with_porteo_header.py <template_path> <output_path> <json_data> <header_template_path>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    output_path = sys.argv[2]
    json_data = sys.argv[3]
    header_template_path = sys.argv[4]
    
    # Parser les données JSON
    data = json.loads(json_data)
    
    fill_template_with_porteo(template_path, output_path, data, header_template_path)

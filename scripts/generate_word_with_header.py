#!/usr/bin/env python3
"""
Script pour générer des documents Word avec en-tête Justicia ou Porteo
"""
import sys
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import numpy as np
from io import BytesIO

def add_justicia_header(doc):
    """Ajoute l'en-tête Justicia au document"""
    logo_path = "/home/ubuntu/papierentetejusticia.docx"
    if os.path.exists(logo_path):
        try:
            # Charger le template et copier TOUT l'en-tête
            template = Document(logo_path)
            template_header = template.sections[0].header
            doc_header = doc.sections[0].header
            
            # Copier tous les éléments de l'en-tête
            for element in template_header._element:
                doc_header._element.append(element)
        except Exception as e:
            print(f"Erreur lors de l'ajout du logo: {e}", file=sys.stderr)
            # Fallback: texte simple
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0]
            run = header_para.add_run("JUSTICIA")
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(128, 0, 128)
    else:
        # Fallback si pas de template
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        run = header_para.add_run("JUSTICIA")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(128, 0, 128)

def add_porteo_header(doc):
    """Ajoute l'en-tête Porteo Group au document"""
    logo_path = "/home/ubuntu/PAPIERENTETEPORTEOGROUP2025.docx"
    if os.path.exists(logo_path):
        try:
            # Charger le template et copier TOUT l'en-tête
            template = Document(logo_path)
            template_header = template.sections[0].header
            doc_header = doc.sections[0].header
            
            # Copier tous les éléments de l'en-tête
            for element in template_header._element:
                doc_header._element.append(element)
        except Exception as e:
            print(f"Erreur lors de l'ajout du logo: {e}", file=sys.stderr)
            # Fallback: texte simple
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0]
            run = header_para.add_run("PORTEO GROUP")
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(218, 165, 32)
    else:
        # Fallback si pas de template
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        run = header_para.add_run("PORTEO GROUP")
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(218, 165, 32)

def generate_risk_chart(risks_data, chart_type="bar"):
    """
    Génère un graphique de risque et retourne le chemin du fichier image
    risks_data: dict avec {"nom_risque": score, ...}
    chart_type: "bar", "heatmap", ou "density"
    """
    plt.figure(figsize=(8, 5))
    
    if chart_type == "bar":
        # Graphique en barres
        names = list(risks_data.keys())
        scores = list(risks_data.values())
        colors = ['#90EE90' if s < 4 else '#FFD700' if s < 7 else '#FF6B6B' for s in scores]
        
        plt.barh(names, scores, color=colors)
        plt.xlabel('Score de Risque')
        plt.xlim(0, 10)
        plt.title('Scores de Risque par Zone')
        plt.grid(axis='x', alpha=0.3)
        
    elif chart_type == "heatmap":
        # Carte de chaleur simplifiée
        fig, ax = plt.subplots(figsize=(8, 2))
        data_array = np.array([list(risks_data.values())])
        im = ax.imshow(data_array, cmap='RdYlGn_r', aspect='auto', vmin=0, vmax=10)
        ax.set_xticks(range(len(risks_data)))
        ax.set_xticklabels(list(risks_data.keys()), rotation=45, ha='right')
        ax.set_yticks([])
        plt.colorbar(im, ax=ax, label='Score de Risque')
        plt.title('Carte de Chaleur des Scores de Risque')
        plt.tight_layout()
    
    # Sauvegarder dans un buffer
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    
    # Sauvegarder temporairement
    temp_path = f"/tmp/risk_chart_{chart_type}.png"
    with open(temp_path, 'wb') as f:
        f.write(buf.read())
    
    return temp_path

def generate_report(data, output_path, header_type="justicia"):
    """Génère un rapport Word avec en-tête"""
    doc = Document()
    
    # Ajouter l'en-tête approprié
    if header_type == "justicia":
        add_justicia_header(doc)
    elif header_type == "porteo":
        add_porteo_header(doc)
    
    # Titre principal
    title = doc.add_heading(data.get("title", "Rapport d'Analyse"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph(f"Date: {data.get('date', '')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Ligne vide
    
    # Contenu
    for section in data.get("sections", []):
        # Titre de section
        doc.add_heading(section.get("title", ""), level=1)
        
        # Contenu de section
        content = section.get("content", "")
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph(content)
        
        doc.add_paragraph()  # Ligne vide
    
    # Ajouter les graphiques de risque si disponibles
    risks_data = data.get("risks", {})
    if risks_data:
        doc.add_heading("Graphiques de Risque", level=1)
        
        # Graphique en barres
        try:
            bar_chart_path = generate_risk_chart(risks_data, "bar")
            doc.add_paragraph("Scores de Risque par Zone (Graphique en Barres)")
            doc.add_picture(bar_chart_path, width=Inches(6))
            doc.add_paragraph()  # Ligne vide
        except Exception as e:
            print(f"Erreur lors de la génération du graphique en barres: {e}", file=sys.stderr)
        
        # Carte de chaleur
        try:
            heatmap_path = generate_risk_chart(risks_data, "heatmap")
            doc.add_paragraph("Carte de Chaleur des Scores de Risque")
            doc.add_picture(heatmap_path, width=Inches(6))
            doc.add_paragraph()  # Ligne vide
        except Exception as e:
            print(f"Erreur lors de la génération de la carte de chaleur: {e}", file=sys.stderr)
    
    # Sauvegarder
    doc.save(output_path)
    print(f"Document généré: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python generate_word_with_header.py <json_file_or_data> <output_path> <header_type>")
        sys.exit(1)
    
    json_input = sys.argv[1]
    output_path = sys.argv[2]
    header_type = sys.argv[3]
    
    try:
        # Vérifier si c'est un fichier ou une chaîne JSON
        if os.path.exists(json_input):
            with open(json_input, 'r', encoding='utf-8') as f:
                data = json.load(f)
        else:
            data = json.loads(json_input)
        
        generate_report(data, output_path, header_type)
    except Exception as e:
        print(f"Erreur: {e}", file=sys.stderr)
        sys.exit(1)

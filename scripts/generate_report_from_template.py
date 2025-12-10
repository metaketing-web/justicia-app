#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pour générer des rapports Word avec parsing markdown
"""
import sys
import json
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def parse_markdown_line(doc, line):
    """Parse une ligne markdown et l'ajoute au document"""
    line = line.strip()
    
    if not line:
        return
    
    # Titre niveau 1 (###)
    if line.startswith('### '):
        title = line.replace('### ', '').strip()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xBA, 0x8A, 0x52)  # Porteo Gold
        p.space_after = Pt(12)
        p.space_before = Pt(18)
        return True
    
    # Titre niveau 2 (####)
    elif line.startswith('#### '):
        title = line.replace('#### ', '').strip()
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xBA, 0x8A, 0x52)  # Porteo Gold
        p.space_after = Pt(10)
        p.space_before = Pt(14)
        return True
    
    # Liste à puces (-)
    elif line.startswith('- '):
        text = line.replace('- ', '').strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run('• ')
        add_formatted_text(p, text)
        p.space_after = Pt(6)
        return True
    
    # Liste numérotée
    elif re.match(r'^\d+\.\s', line):
        match = re.match(r'^(\d+)\.\s', line)
        num = match.group(1)
        text = re.sub(r'^\d+\.\s', '', line).strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f'{num}. ')
        add_formatted_text(p, text)
        p.space_after = Pt(6)
        return True
    
    return False

def add_formatted_text(paragraph, text):
    """Ajoute du texte avec formatage markdown (gras, italique)"""
    # Remplacer **texte** par gras
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        else:
            run = paragraph.add_run(part)
            run.font.color.rgb = RGBColor(0x17, 0x23, 0x2E)  # Porteo Dark Blue

def parse_markdown_content(doc, content):
    """Parse le contenu markdown et l'ajoute au document"""
    # Nettoyer le contenu
    content = content.replace('\\n', '\n')
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Essayer de parser comme titre ou liste
        if parse_markdown_line(doc, line):
            i += 1
            continue
        
        # Sinon, accumuler les lignes du paragraphe
        if line:
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line or next_line.startswith(('#', '-')) or re.match(r'^\d+\.\s', next_line):
                    break
                para_lines.append(next_line)
                i += 1
            
            # Créer le paragraphe
            text = ' '.join(para_lines)
            p = doc.add_paragraph()
            add_formatted_text(p, text)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p.space_after = Pt(12)
        else:
            i += 1

def generate_report_from_template(json_path, output_path, template_type='justicia'):
    """Génère un rapport Word à partir du modèle"""
    
    # Charger les données JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Charger le template
    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(script_dir, f'template_{template_type}.docx')
    
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
    
    # Ajouter un espace après l'en-tête
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Titre principal
    title = data.get('title', 'Rapport')
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_title = p_title.add_run(title)
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0xBA, 0x8A, 0x52)  # Porteo Gold
    p_title.space_after = Pt(6)
    
    # Date
    date = data.get('date', '')
    if date:
        p_date = doc.add_paragraph()
        p_date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_date = p_date.add_run(f"Date: {date}")
        run_date.font.size = Pt(11)
        run_date.font.italic = True
        run_date.font.color.rgb = RGBColor(128, 128, 128)
        p_date.space_after = Pt(24)
    
    # Ligne de séparation
    doc.add_paragraph('_' * 80)
    
    # Vérifier si c'est un rapport d'analyse (analysisResults) ou un document standard (sections)
    if 'plainLanguageSummary' in data or 'flags' in data or 'riskAssessment' in data:
        # Format analysisResults
        
        # Résumé
        if 'plainLanguageSummary' in data:
            p_section = doc.add_paragraph()
            run_section = p_section.add_run('Résumé')
            run_section.font.size = Pt(18)
            run_section.font.bold = True
            run_section.font.color.rgb = RGBColor(75, 0, 130)
            p_section.space_before = Pt(18)
            p_section.space_after = Pt(12)
            parse_markdown_content(doc, data['plainLanguageSummary'])
        
        # Évaluation des risques
        if 'riskAssessment' in data:
            risk = data['riskAssessment']
            p_section = doc.add_paragraph()
            run_section = p_section.add_run('Évaluation des Risques')
            run_section.font.size = Pt(18)
            run_section.font.bold = True
            run_section.font.color.rgb = RGBColor(75, 0, 130)
            p_section.space_before = Pt(18)
            p_section.space_after = Pt(12)
            
            if 'overallSummary' in risk:
                parse_markdown_content(doc, risk['overallSummary'])
            
            if 'risks' in risk:
                for r in risk['risks']:
                    text = f"- **{r.get('area', '')}** : {r.get('assessment', '')} (Score: {r.get('score', 0)}/10)"
                    p = doc.add_paragraph()
                    add_formatted_text(p, text)
                    p.space_after = Pt(6)
        
        # Analyses IA
        if 'aiInsights' in data:
            insights = data['aiInsights']
            p_section = doc.add_paragraph()
            run_section = p_section.add_run('Analyses IA')
            run_section.font.size = Pt(18)
            run_section.font.bold = True
            run_section.font.color.rgb = RGBColor(75, 0, 130)
            p_section.space_before = Pt(18)
            p_section.space_after = Pt(12)
            
            if isinstance(insights, dict) and 'overallSummary' in insights:
                parse_markdown_content(doc, insights['overallSummary'])
                if 'recommendations' in insights:
                    for rec in insights['recommendations']:
                        text = f"- **{rec.get('recommendation', '')}** : {rec.get('justification', '')}"
                        p = doc.add_paragraph()
                        add_formatted_text(p, text)
                        p.space_after = Pt(6)
            elif isinstance(insights, str):
                parse_markdown_content(doc, insights)
        
        # Signalements
        if 'flags' in data and len(data['flags']) > 0:
            p_section = doc.add_paragraph()
            run_section = p_section.add_run('Signalements Trouvés')
            run_section.font.size = Pt(18)
            run_section.font.bold = True
            run_section.font.color.rgb = RGBColor(75, 0, 130)
            p_section.space_before = Pt(18)
            p_section.space_after = Pt(12)
            
            for flag in data['flags']:
                # Titre du flag
                p = doc.add_paragraph()
                run = p.add_run(flag.get('title', ''))
                run.font.size = Pt(14)
                run.font.bold = True
                p.space_before = Pt(12)
                p.space_after = Pt(6)
                
                # Clause
                if 'clause' in flag:
                    p = doc.add_paragraph()
                    add_formatted_text(p, f"**Clause** : {flag['clause']}")
                    p.space_after = Pt(6)
                
                # Explication
                if 'explanation' in flag:
                    p = doc.add_paragraph()
                    add_formatted_text(p, f"**Explication** : {flag['explanation']}")
                    p.space_after = Pt(6)
                
                # Réécriture suggérée
                if 'suggestedRewrite' in flag and flag['suggestedRewrite']:
                    p = doc.add_paragraph()
                    add_formatted_text(p, f"**Réécriture suggérée** : {flag['suggestedRewrite']}")
                    p.space_after = Pt(12)
    
    else:
        # Format sections standard
        sections = data.get('sections', [])
        for section in sections:
            section_title = section.get('title', '')
            section_content = section.get('content', '')
            
            # Titre de section
            if section_title:
                p_section = doc.add_paragraph()
                run_section = p_section.add_run(section_title)
                run_section.font.size = Pt(18)
                run_section.font.bold = True
                run_section.font.color.rgb = RGBColor(75, 0, 130)  # Indigo
                p_section.space_before = Pt(18)
                p_section.space_after = Pt(12)
            
            # Contenu (parser markdown)
            if section_content:
                parse_markdown_content(doc, section_content)
    
    # Sauvegarder
    doc.save(output_path)
    print(f"[TEMPLATE] Document généré: {output_path}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python3 generate_report_from_template.py <json_path> <output_path> [template_type]")
        sys.exit(1)
    
    json_path = sys.argv[1]
    output_path = sys.argv[2]
    template_type = sys.argv[3] if len(sys.argv) > 3 else 'justicia'
    
    generate_report_from_template(json_path, output_path, template_type)

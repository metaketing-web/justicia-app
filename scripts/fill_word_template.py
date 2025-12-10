#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pour remplir un template Word avec des données
"""
import sys
import json
import os
import re
from docx import Document

def replace_placeholder(paragraph, placeholder, value):
    """Remplace un placeholder dans un paragraphe"""
    if placeholder in paragraph.text:
        # Remplacer dans les runs
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(value))

def fill_template(template_path, output_path, data):
    """Remplit un template Word avec les données fournies"""
    
    # Charger le template
    doc = Document(template_path)
    
    # Créer les placeholders à partir des données
    placeholders = {}
    for key, value in data.items():
        # Format: {{nom_champ}}
        placeholder = f"{{{{{key}}}}}"
        placeholders[placeholder] = value
    
    # Remplacer dans tous les paragraphes
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            replace_placeholder(paragraph, placeholder, value)
    
    # Remplacer dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in placeholders.items():
                        replace_placeholder(paragraph, placeholder, value)
    
    # Remplacer dans les en-têtes et pieds de page
    for section in doc.sections:
        # En-tête
        header = section.header
        for paragraph in header.paragraphs:
            for placeholder, value in placeholders.items():
                replace_placeholder(paragraph, placeholder, value)
        
        # Pied de page
        footer = section.footer
        for paragraph in footer.paragraphs:
            for placeholder, value in placeholders.items():
                replace_placeholder(paragraph, placeholder, value)
    
    # Sauvegarder
    doc.save(output_path)
    print(f"[TEMPLATE] Document généré: {output_path}")

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("Usage: python3 fill_word_template.py <template_path> <output_path> <json_data>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    output_path = sys.argv[2]
    json_data = sys.argv[3]
    
    # Parser les données JSON
    data = json.loads(json_data)
    
    fill_template(template_path, output_path, data)
